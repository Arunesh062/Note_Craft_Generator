import os
import re
from datetime import datetime

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

OUTPUTS_DIR = os.path.join(os.path.dirname(__file__), "..", "outputs")


# ─────────────────────────────────────────────
# MAIN EXPORT FUNCTION
# ─────────────────────────────────────────────
def export_documents(mom_json: dict, session_id: str) -> tuple:

    os.makedirs(OUTPUTS_DIR, exist_ok=True)

    raw_title = (
        mom_json.get("session_title")
        or mom_json.get("title")
        or "MoM_Report"
    )

    clean_name = re.sub(r"[^\w\s-]", "", str(raw_title))
    clean_name = re.sub(r"\s+", "_", clean_name.strip())
    clean_name = clean_name[:60]

    filename = f"{clean_name}_{session_id[:8]}"

    docx_path = os.path.join(
        OUTPUTS_DIR,
        f"{filename}.docx"
    )

    _generate_docx(mom_json, docx_path)

    return None, f"/outputs/{filename}.docx"


# ─────────────────────────────────────────────
# DOCX GENERATION
# ─────────────────────────────────────────────
def _generate_docx(data: dict, path: str):

    doc = Document()

    # PAGE SETTINGS
    section = doc.sections[0]

    section.top_margin = Pt(40)
    section.bottom_margin = Pt(40)
    section.left_margin = Pt(40)
    section.right_margin = Pt(40)

    # FOOTER
    footer = section.footer
    footer_para = footer.paragraphs[0]
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    _add_page_number(footer_para)

    # ─────────────────────────────────────────
    # HEADER TEXT
    # ─────────────────────────────────────────
    dept_para = doc.add_paragraph()
    dept_para.alignment = WD_ALIGN_PARAGRAPH.LEFT

    run = dept_para.add_run(
        "Department : CSE\nNandha Engineering College"
    )

    run.bold = True
    run.font.size = Pt(10)

    # ─────────────────────────────────────────
    # HEADER TABLE
    # ─────────────────────────────────────────
    table = doc.add_table(rows=3, cols=6)
    table.style = "Table Grid"

    def set_cell(r, c, text, bold=False, center=False):

        cell = table.cell(r, c)
        cell.text = str(text)

        para = cell.paragraphs[0]

        if center:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        run = para.runs[0]
        run.font.size = Pt(9)

        if bold:
            run.bold = True

    # ROW 1
    set_cell(0, 0, "Name of the\nMeeting", True)

    meeting_title = (
        data.get("session_title")
        or data.get("title")
        or "Faculty Meeting"
    )

    meeting_cell = table.cell(0, 1).merge(table.cell(0, 2))
    meeting_cell.text = str(meeting_title)
    meeting_cell.paragraphs[0].runs[0].font.size = Pt(9)

    set_cell(0, 3, "Number", True)

    num_cell = table.cell(0, 4).merge(table.cell(0, 5))
    num_cell.text = f"2025-26/ {datetime.now().strftime('%m')}"
    num_cell.paragraphs[0].runs[0].font.size = Pt(9)

    # ROW 2
    set_cell(1, 0, "Date", True)

    set_cell(
        1,
        1,
        datetime.now().strftime("%Y-%m-%d")
    )

    set_cell(1, 2, "Time", True)

    set_cell(
        1,
        3,
        data.get("time") or "None"
    )

    set_cell(1, 4, "Venue", True)

    set_cell(
        1,
        5,
        data.get("venue")
        or "Block III – Staff Floor Cabin"
    )

    # ROW 3
    set_cell(2, 0, "Members Present", True)

    members_cell = table.cell(2, 1).merge(
        table.cell(2, 5)
    )

    participants = data.get("participants") or []

    if isinstance(participants, list) and participants:

        members_cell.text = ", ".join(
            [str(p) for p in participants if p]
        )

    else:
        members_cell.text = "HoD and All faculty members"

    members_cell.paragraphs[0].runs[0].font.size = Pt(9)

    doc.add_paragraph()

    # ─────────────────────────────────────────
    # TITLE
    # ─────────────────────────────────────────
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    run = title_para.add_run(
        "Minutes of the Meeting"
    )

    run.bold = True
    run.font.size = Pt(12)
    run.underline = True

    doc.add_paragraph()

    # ─────────────────────────────────────────
    # MAIN TABLE
    # ─────────────────────────────────────────
    mom_table = doc.add_table(rows=1, cols=4)
    mom_table.style = "Table Grid"

    headers = [
        "Category",
        "Points Discussed",
        "Responsibility",
        "Target Date"
    ]

    hdr_cells = mom_table.rows[0].cells

    for i, h in enumerate(headers):

        hdr_cells[i].text = h

        para = hdr_cells[i].paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        run = para.runs[0]
        run.bold = True
        run.font.size = Pt(10)

        _set_cell_background(
            hdr_cells[i],
            "D9D9D9"
        )

    # ─────────────────────────────────────────
    # CATEGORY SECTION
    # ─────────────────────────────────────────
    categories = data.get("categories") or []

    # AUTO CONVERT
    if not categories:

        fallback_points = (
            data.get("key_takeaways")
            or data.get("information_items")
            or []
        )

        if not isinstance(fallback_points, list):
            fallback_points = [str(fallback_points)]

        categories = [
            {
                "name": "Meeting Discussion",
                "points": fallback_points,
                "responsibility": "All",
                "target_date": "Continuous"
            }
        ]

    # CREATE ROWS
    for idx, cat in enumerate(categories, start=1):

        if not isinstance(cat, dict):
            continue

        row = mom_table.add_row()

        # CATEGORY
        category_cell = row.cells[0]

        category_name = (
            f"{idx}. "
            f"{cat.get('name') or 'General'}"
        )

        category_cell.text = category_name

        cat_run = (
            category_cell.paragraphs[0]
            .runs[0]
        )

        cat_run.bold = True
        cat_run.font.size = Pt(9)

        # POINTS
        points_cell = row.cells[1]

        points = cat.get("points") or []

        if not isinstance(points, list):
            points = [str(points)]

        if points:

            first_para = points_cell.paragraphs[0]
            first_para.style = "List Bullet"

            first_run = first_para.add_run(
                str(points[0])
            )

            first_run.font.size = Pt(9)

            for p_text in points[1:]:

                p = points_cell.add_paragraph(
                    style="List Bullet"
                )

                r = p.add_run(str(p_text))
                r.font.size = Pt(9)

        else:

            points_cell.text = (
                "Discussion conducted."
            )

        # RESPONSIBILITY
        res_cell = row.cells[2]

        res_cell.text = str(
            cat.get("responsibility")
            or "All"
        )

        res_para = res_cell.paragraphs[0]

        res_para.alignment = (
            WD_ALIGN_PARAGRAPH.CENTER
        )

        res_para.runs[0].font.size = Pt(9)

        # TARGET DATE
        td_cell = row.cells[3]

        td_cell.text = str(
            cat.get("target_date")
            or "Continuous"
        )

        td_para = td_cell.paragraphs[0]

        td_para.alignment = (
            WD_ALIGN_PARAGRAPH.CENTER
        )

        td_para.runs[0].font.size = Pt(9)

    doc.add_paragraph()

    # ─────────────────────────────────────────
    # INFORMATION ITEMS
    # ─────────────────────────────────────────
    info_para = doc.add_paragraph()

    run = info_para.add_run(
        "Information Items"
    )

    run.bold = True
    run.font.size = Pt(11)
    run.underline = True

    info_items = (
        data.get("information_items")
        or data.get("key_takeaways")
        or []
    )

    if not isinstance(info_items, list):
        info_items = [str(info_items)]

    if not info_items:

        info_items = [
            "Faculty members are requested to complete pending work.",
            "Upcoming activities will be communicated shortly.",
            "All members should follow department schedule."
        ]

    for item in info_items:

        p = doc.add_paragraph(
            style="List Bullet"
        )

        r = p.add_run(str(item))
        r.font.size = Pt(9)

    # ─────────────────────────────────────────
    # SAVE
    # ─────────────────────────────────────────
    try:

        doc.save(path)

        print("✅ DOCX saved:", path)

    except Exception as e:

        print("❌ DOCX save failed:", e)

        raise


# ─────────────────────────────────────────────
# CELL BACKGROUND
# ─────────────────────────────────────────────
def _set_cell_background(cell, color):

    tc_pr = cell._tc.get_or_add_tcPr()

    shd = OxmlElement("w:shd")

    shd.set(qn("w:fill"), color)

    tc_pr.append(shd)


# ─────────────────────────────────────────────
# PAGE NUMBER
# ─────────────────────────────────────────────
def _add_page_number(paragraph):

    paragraph.add_run("Page ")

    # PAGE
    run = paragraph.add_run()

    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(
        qn("w:fldCharType"),
        "begin"
    )

    run._r.append(fld_begin)

    run = paragraph.add_run()

    instr = OxmlElement("w:instrText")
    instr.set(
        qn("xml:space"),
        "preserve"
    )

    instr.text = "PAGE"

    run._r.append(instr)

    run = paragraph.add_run()

    fld_end = OxmlElement("w:fldChar")

    fld_end.set(
        qn("w:fldCharType"),
        "end"
    )

    run._r.append(fld_end)

    paragraph.add_run(" of ")

    # NUMPAGES
    run = paragraph.add_run()

    fld_begin2 = OxmlElement("w:fldChar")

    fld_begin2.set(
        qn("w:fldCharType"),
        "begin"
    )

    run._r.append(fld_begin2)

    run = paragraph.add_run()

    instr2 = OxmlElement("w:instrText")

    instr2.set(
        qn("xml:space"),
        "preserve"
    )

    instr2.text = "NUMPAGES"

    run._r.append(instr2)

    run = paragraph.add_run()

    fld_end2 = OxmlElement("w:fldChar")

    fld_end2.set(
        qn("w:fldCharType"),
        "end"
    )

    run._r.append(fld_end2)